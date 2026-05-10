"""Regenerate the Chances Assessment doc with properly extracted content."""
import json, re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/home/alba/atticus-harness-v2/matters/napier-accommodation-arrears/_artifacts/case-report-1778253473719.json"
TARGET = "/home/alba/Documents/Anfal + Mom - Fighting Chance/06-Chances-Assessment.docx"

# Read and extract the actual content
with open(SRC) as f:
    d = json.load(f)

raw = d["content"]

# Strip the inner JSON wrapper: find "content": " and extract
start_marker = '"content": "'
idx = raw.find(start_marker)
if idx > 0:
    start = idx + len(start_marker)
    text = raw[start:].rstrip()
    if text.endswith('"'):
        text = text[:-1]
    # Unescape
    text = text.replace('\\n', '\n')
    text = text.replace('\\"', '"')
    text = text.replace('\\\\', '\\')
    text = re.sub(r'\\u[0-9a-fA-F]{4}', lambda m: chr(int(m.group(0)[2:], 16)), text)
    # Now strip the outer JSON title/type lines
    # Remove the ### heading that appears after the title wrapper
    lines = text.split('\n')
    # Find where the actual markdown content starts (after ## Chances Assessment...)
    content_lines = []
    found_start = False
    for line in lines:
        if line.startswith('## Chances Assessment'):
            found_start = True
        if found_start:
            content_lines.append(line)
    if content_lines:
        text = '\n'.join(content_lines)
else:
    text = raw

print(f"Clean content: {len(text)} chars")

# Build the .docx
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Chances Assessment & Escalation Roadmap")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anfal Elbushra (40805511) - Edinburgh Napier University Accommodation Arrears")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2500" * 60)
run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
run.font.size = Pt(8)

# Process each line
for line in text.split('\n'):
    stripped = line.strip()
    
    # Skip empty lines
    if not stripped:
        continue
    
    # Divider lines (---)
    if re.match(r'^-{3,}$', stripped):
        p = doc.add_paragraph()
        run = p.add_run("\u2500" * 60)
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        run.font.size = Pt(6)
        continue
    
    # Headings (### or ##)
    heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        h_text = heading_match.group(2)
        if level <= 2:
            h = doc.add_heading(h_text, level=1)
        else:
            h = doc.add_heading(h_text, level=2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        continue
    
    # Bold section headers like **Assessment: Moderate chance**
    if stripped.startswith('**') and stripped.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped[2:-2])
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        continue
    
    # Normal paragraph with inline **bold** markers
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    parts = re.split(r'(\*\*.*?\*\*)', stripped)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        else:
            if part.strip():
                run = p.add_run(part)
                run.font.size = Pt(11)
    
    # Indent bullet points
    if stripped.startswith('- ') or stripped.startswith('* '):
        p.paragraph_format.left_indent = Cm(1.0)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepare-only advocacy materials - Anfal Elbushra (40805511)")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.save(TARGET)
print(f"Saved: {TARGET} ({os.path.getsize(TARGET)} bytes)")
