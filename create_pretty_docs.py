"""Convert legal content to properly formatted .docx files."""
import json, re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

TARGET_DIR = "/home/alba/Documents/Anfal + Mom - Fighting Chance"
HARNESS = "/home/alba/atticus-harness-v2"

def read_artifact(matter, artifact_file):
    """Read an artifact JSON and return its content as plain text."""
    path = f"{HARNESS}/matters/{matter}/_artifacts/{artifact_file}"
    with open(path) as f:
        d = json.load(f)
    content = d.get("content", "")
    # Handle double-encoded JSON (for case-report artifacts that contain a JSON document)
    if content.startswith("{"):
        try:
            inner = json.loads(content)
            if isinstance(inner, dict) and "content" in inner:
                content = inner["content"]
        except json.JSONDecodeError:
            pass  # use raw content
    return content

def make_docx(title, subtitle, body_text, filename):
    """Create a nicely formatted .docx from body_text."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
    run.font.size = Pt(8)
    
    # Process body text line by line for formatting
    prev_bold = False
    for line in body_text.split("\n"):
        stripped = line.strip()
        
        # Skip divider lines
        if re.match(r'^[=\-]{5,}$', stripped) or re.match(r'^[-─]{5,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped[:60])
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
            run.font.size = Pt(6)
            continue
        
        # Empty line = paragraph break
        if not stripped:
            # Add a small spacer instead of full empty para
            continue
        
        # Detect markdown headings (## or ###)
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level <= 2:
                h = doc.add_heading(text, level=1)
            else:
                h = doc.add_heading(text, level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            continue
        
        # Detect numbered sections like "1. Title" or "**1. Title**"
        section_match = re.match(r'^\*{0,2}(\d+)\.\s+(.+?)\*{0,2}$', stripped)
        if section_match and len(stripped) < 80:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            continue
        
        # Normal paragraph with inline formatting
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        # Split by ** markers for bold
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                text = part[2:-2]
                if text.strip():
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(11)
            else:
                if part.strip():
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        
        # Handle bullet list items
        if stripped.startswith("- ") or stripped.startswith("* "):
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Anfal Elbushra (40805511) — Prepare-only advocacy materials")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    
    outpath = os.path.join(TARGET_DIR, filename)
    doc.save(outpath)
    size = os.path.getsize(outpath)
    print(f"  ✓ {filename} ({size:,} bytes)")
    return outpath


def extract_text_from_deliverable(content, deliverable_num):
    """Extract a specific deliverable section from the master 6-in-1 document."""
    pattern = f"DELIVERABLE {deliverable_num}:.*?\\n=+\\n\\n(.*?)(?=\\n\\n=+\\n\\nDELIVERABLE {deliverable_num+1}:|\\Z)"
    match = re.search(pattern, content, re.DOTALL)
    if match:
        return match.group(1).strip()
    return None


# ─── Document specs ──────────────────────────────────────────

docs = []

# 1. COMPLAINT LETTER (v2 from codex55 - latest)
complaint = read_artifact("anfal-napier-codex55-fixed", "case-draft-1778364356580.json")
docs.append(("Formal Complaint Letter", "Edinburgh Napier University — Anfal Elbushra (40805511)", complaint, "01-Formal-Complaint-Letter.docx"))

# 2. GUARANTOR PURSUIT GUIDE (from codex55 - latest, most detailed)
guarantor = read_artifact("anfal-napier-codex55-fixed", "case-report-1778364466892.json")
docs.append(("Guarantor Pursuit Guide", "Protecting Leila Hassan — Devil's Advocate Preparation", guarantor, "02-Guarantor-Pursuit-Guide.docx"))

# 3. ENTITLEMENTS REPORT (from codex55 - unique)
entitlements = read_artifact("anfal-napier-codex55-fixed", "case-report-1778364478729.json")
docs.append(("Entitlements Report", "What Anfal is Owed — Edinburgh Napier's Failures", entitlements, "03-Entitlements-Report.docx"))

# 4. MASTER ACTION PLAN v2 (from codex55 - latest)
action_plan = read_artifact("anfal-napier-codex55-fixed", "case-task-1778364502704.json")
docs.append(("Master Action Plan", "Tomorrow Morning Onwards — Day-by-Day Strategy", action_plan, "04-Master-Action-Plan.docx"))

# 5. LEGAL POSITION PAPER (from napier master doc - unique content)
napier_master = read_artifact("napier-accommodation-arrears", "case-task-1778364915529.json")
legal_pos = extract_text_from_deliverable(napier_master, 3)
if not legal_pos:
    legal_pos = "Legal Position Paper content could not be extracted."
docs.append(("Legal Position Paper", "Anfal Elbushra v Edinburgh Napier — Legal Analysis", legal_pos, "05-Legal-Position-Paper.docx"))

# 6. CHANCES ASSESSMENT (from napier - unique content)
chances_raw = read_artifact("napier-accommodation-arrears", "case-report-1778253473719.json")
# If it starts with {, it's JSON with escaped content, try to parse the inner
if chances_raw.startswith("{"):
    try:
        inner = json.loads(chances_raw)
        if isinstance(inner, dict):
            chances_raw = inner.get("content", chances_raw)
    except:
        pass
docs.append(("Chances Assessment & Escalation Roadmap", "Likelihood of Each Outcome — Step-by-Step Escalation", chances_raw, "06-Chances-Assessment.docx"))


# ─── Generate ──────────────────────────────────────────
print("Creating pretty .docx files...")
os.makedirs(TARGET_DIR, exist_ok=True)

paths = []
for title, subtitle, content, filename in docs:
    path = make_docx(title, subtitle, content, filename)
    paths.append(path)

print(f"\nAll done. {len(paths)} files in {TARGET_DIR}")
